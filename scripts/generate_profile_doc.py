"""
Generate a professional Profcaria Company Profile document (DOCX).
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# Custom heading styles
for level, (size, color) in enumerate([
    (Pt(28), RGBColor(0x0A, 0x0A, 0x0A)),  # Heading 1
    (Pt(18), RGBColor(0x1A, 0x1A, 0x1A)),  # Heading 2
    (Pt(14), RGBColor(0x2A, 0x2A, 0x2A)),  # Heading 3
], start=1):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.size = size
    h.font.bold = True
    h.font.color.rgb = color
    h.paragraph_format.space_before = Pt(18 if level > 1 else 6)
    h.paragraph_format.space_after = Pt(8)


def add_divider(doc):
    """Add a thin horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/></w:pBdr>')
    pPr.append(pBdr)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a clean, professional table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A1A1A"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Alternating row colors
            if r_idx % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    # Column widths
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def add_bullet(doc, text, bold_prefix=None):
    """Add a bullet point with optional bold prefix."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run = p.add_run(f" — {text}")
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'


def add_body(doc, text):
    """Add body paragraph."""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    return p


# ═══════════════════════════════════════════════════════════
# COVER / TITLE
# ═══════════════════════════════════════════════════════════

# Spacer
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('PROFCARIA')
run.bold = True
run.font.size = Pt(48)
run.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)
run.font.name = 'Calibri'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('The Professional Network for the Modern Era')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = 'Calibri'
run.italic = True

doc.add_paragraph()

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run('Company Profile & Platform Overview')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
run.font.name = 'Calibri'

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('www.profcaria.com')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
run.font.name = 'Calibri'
run.bold = True

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('February 2026')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = 'Calibri'

conf = doc.add_paragraph()
conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = conf.add_run('CONFIDENTIAL')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
run.font.name = 'Calibri'
run.bold = True

# Page break
doc.add_page_break()


# ═══════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════

doc.add_heading('Table of Contents', level=1)
add_divider(doc)

toc_items = [
    '1. About Profcaria',
    '2. Our Mission',
    '3. Our Vision',
    '4. Our Mandate & Goals',
    '5. The Platform — What We Do',
    '6. Key Features',
    '7. The Profcaria Ecosystem',
    '8. For Employers — Hiring Made Smarter',
    '9. For Professionals — Your Career, Elevated',
    '10. Pricing Plans',
    '11. How We Compare',
    '12. No Ads. Ever.',
    '13. Security & Trust',
    '14. Why Profcaria',
    '15. Contact Information',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════
# 1. ABOUT PROFCARIA
# ═══════════════════════════════════════════════════════════

doc.add_heading('1. About Profcaria', level=1)
add_divider(doc)

add_body(doc,
    "Profcaria is a next-generation professional networking and talent acquisition platform "
    "designed for the modern workforce. Built from the ground up with Africa's unique professional "
    "landscape in mind, Profcaria bridges the gap between exceptional talent and forward-thinking "
    "employers through intelligent technology, intuitive design, and a deep commitment to "
    "creating meaningful professional connections."
)

add_body(doc,
    "Unlike traditional job boards or social media platforms repurposed for hiring, Profcaria is "
    "purpose-built as a unified ecosystem where professionals grow their careers, showcase their "
    "work, and connect with opportunities — while employers gain access to powerful AI-driven tools "
    "that make hiring faster, smarter, and more cost-effective."
)

add_body(doc,
    "Profcaria is more than a tool — it is a professional community. A space where ambition finds "
    "its home."
)


# ═══════════════════════════════════════════════════════════
# 2. OUR MISSION
# ═══════════════════════════════════════════════════════════

doc.add_heading('2. Our Mission', level=1)
add_divider(doc)

mission = doc.add_paragraph()
mission.alignment = WD_ALIGN_PARAGRAPH.CENTER
mission.paragraph_format.space_before = Pt(12)
mission.paragraph_format.space_after = Pt(12)
run = mission.add_run(
    '"To democratise access to professional opportunities and empower '
    'every individual — regardless of background, degree, or network — '
    'to build a career that reflects their true potential."'
)
run.italic = True
run.font.size = Pt(13)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

add_body(doc,
    "We believe that talent is evenly distributed, but opportunity is not. Profcaria exists to "
    "close that gap — connecting skilled professionals with the right employers through "
    "technology that sees beyond CVs and credentials to find genuine fit."
)


# ═══════════════════════════════════════════════════════════
# 3. OUR VISION
# ═══════════════════════════════════════════════════════════

doc.add_heading('3. Our Vision', level=1)
add_divider(doc)

vision = doc.add_paragraph()
vision.alignment = WD_ALIGN_PARAGRAPH.CENTER
vision.paragraph_format.space_before = Pt(12)
vision.paragraph_format.space_after = Pt(12)
run = vision.add_run(
    '"To become Africa\'s most trusted and innovative professional platform — '
    'where every career move is informed, every hire is intentional, and every '
    'professional has a seat at the table."'
)
run.italic = True
run.font.size = Pt(13)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

add_body(doc,
    "We envision a future where hiring is driven by skills, not just degrees. Where professionals "
    "are valued for what they can do, not where they went to school. And where technology serves "
    "as an equaliser — giving small companies the same hiring intelligence as global corporations."
)


# ═══════════════════════════════════════════════════════════
# 4. OUR MANDATE & GOALS
# ═══════════════════════════════════════════════════════════

doc.add_heading('4. Our Mandate & Goals', level=1)
add_divider(doc)

add_body(doc, "Profcaria operates with a clear set of goals that guide every decision we make:")

goals = [
    ("Skills-First Hiring", "Shift the hiring paradigm from credential-based to competency-based recruitment, where portfolios and demonstrated skills carry as much weight as formal education."),
    ("Accessible Technology", "Build enterprise-grade hiring tools and make them accessible at price points that work for African businesses — from startups to multinationals."),
    ("Trust & Verification", "Create a verified professional ecosystem where employers can trust candidate profiles and professionals can trust employer legitimacy."),
    ("Data-Driven Decisions", "Provide employers with rich analytics and AI-powered insights so every hiring decision is backed by data, not just gut feeling."),
    ("Community Over Competition", "Foster a professional community where knowledge sharing, mentorship, and genuine connections take precedence over vanity metrics."),
    ("Zero-Noise Experience", "Maintain a clean, ad-free platform that respects users' time and attention. No distractions, no spam, no irrelevant content."),
]
for title, desc in goals:
    add_bullet(doc, desc, bold_prefix=title)


# ═══════════════════════════════════════════════════════════
# 5. THE PLATFORM
# ═══════════════════════════════════════════════════════════

doc.add_heading('5. The Platform — What We Do', level=1)
add_divider(doc)

add_body(doc,
    "Profcaria is a full-stack professional platform that unites networking, hiring, and career "
    "growth into a single, seamless experience. Here is what happens on Profcaria:"
)

doc.add_heading('For Professionals', level=2)
add_bullet(doc, "Create a rich professional profile — showcase your experience, skills, education, portfolio, and references all in one place.")
add_bullet(doc, "Discover and apply to jobs that match your skills and career goals.")
add_bullet(doc, "Build a professional network through meaningful connections — not random follower counts.")
add_bullet(doc, "Share insights, articles, and professional updates through a curated feed.")
add_bullet(doc, "Join and participate in professional communities relevant to your industry.")
add_bullet(doc, "Get verified — earn verification badges that signal credibility to employers.")
add_bullet(doc, "Receive real-time notifications on application status: shortlisted, employed, or updates from your network.")

doc.add_heading('For Employers', level=2)
add_bullet(doc, "Post jobs and reach a pool of verified, skills-assessed candidates.")
add_bullet(doc, "Leverage AI Top Match to instantly surface the best-fit candidates for each role — no manual sifting required.")
add_bullet(doc, "Manage your full hiring pipeline: review applications, shortlist candidates, and make offers — all from one dashboard.")
add_bullet(doc, "Access hiring analytics: pipeline funnel, time-to-fill, conversion rates, and candidate quality metrics.")
add_bullet(doc, "Post to your company feed to attract talent through employer branding.")
add_bullet(doc, "Use location-restricted job listings to target candidates in specific regions or counties.")
add_bullet(doc, "Request and verify professional references directly on-platform.")


# ═══════════════════════════════════════════════════════════
# 6. KEY FEATURES
# ═══════════════════════════════════════════════════════════

doc.add_heading('6. Key Features', level=1)
add_divider(doc)

features = [
    ("AI Top Match", 
     "Our proprietary matching algorithm analyses job requirements against candidate profiles — "
     "including skills, experience, location, and career preferences — to surface the most relevant "
     "candidates for each position. This saves HR teams hours of manual screening and significantly "
     "improves hiring quality."),
    
    ("Skills-First Hiring Toggle",
     "Employers can enable Skills-First mode on any job posting, signalling to candidates that "
     "portfolios and demonstrable skills are prioritised over traditional degree requirements. "
     "This unlocks access to talented professionals who might otherwise be filtered out."),
    
    ("Verification Badges",
     "Professionals earn verification badges (Gray, Blue, or Gold) that appear on their profile "
     "and in search results. Badges signal credibility and boost visibility, giving verified "
     "professionals up to 8x more visibility in employer searches."),
    
    ("Hiring Analytics Dashboard",
     "A comprehensive analytics suite for employers, featuring pipeline funnels, application trends, "
     "conversion rates, job performance metrics, and weekly growth indicators. All data is presented "
     "in real-time, interactive charts."),
    
    ("Professional Feed & Communities",
     "A curated content feed where professionals share articles, insights, and updates — along with "
     "dedicated community spaces for industry-specific discussions. The feed is algorithmically "
     "ranked to prioritise relevant, high-quality content."),
    
    ("Location-Restricted Jobs",
     "Employers on Pro and Enterprise plans can restrict job visibility to candidates in specific "
     "geographic regions. Ideal for branch offices, field roles, or county-specific recruitment."),
    
    ("Reference System",
     "Professionals can request references from previous employers directly through the platform. "
     "References are verified and displayed on profiles, adding a layer of trust to every candidate."),
    
    ("Real-Time Notifications",
     "Both employers and professionals receive instant in-app notifications for every important event: "
     "new applications, shortlisting decisions, messages, connection requests, and community activity."),
    
    ("Dark & Light Mode",
     "The entire platform is designed with both dark and light themes — optimised for comfortable "
     "viewing in any environment, any time of day."),
    
    ("Mobile-First Design",
     "Every screen on Profcaria is fully responsive and optimised for mobile devices, ensuring "
     "a seamless experience whether you are on a phone, tablet, or desktop."),
]

for title, desc in features:
    doc.add_heading(title, level=3)
    add_body(doc, desc)


# ═══════════════════════════════════════════════════════════
# 7. THE ECOSYSTEM
# ═══════════════════════════════════════════════════════════

doc.add_heading('7. The Profcaria Ecosystem', level=1)
add_divider(doc)

add_body(doc,
    "Profcaria is not a single tool — it is an interconnected ecosystem designed to serve every "
    "participant in the professional journey:"
)

eco = [
    ("Professionals", "Build profiles, find jobs, grow networks, join communities, earn badges, get references."),
    ("Employers & HR Teams", "Post jobs, use AI matching, manage pipelines, analyse hiring data, request references."),
    ("Communities", "Industry-specific spaces for discussion, knowledge sharing, and professional networking."),
    ("Content & Feed", "A professional publishing space — share insights, link articles, showcase projects."),
    ("Verification Layer", "Badge system that creates trust across the entire ecosystem. Employers trust candidates; professionals trust employers."),
    ("Payments & Billing", "Transparent, self-serve billing with plan upgrades, one-time options, and prorated switching. Powered by Paystack for secure, local payment processing."),
]

for title, desc in eco:
    add_bullet(doc, desc, bold_prefix=title)

add_body(doc,
    "Every component feeds into the next. When a professional earns a badge, they become more "
    "visible in job searches. When an employer posts a job, AI matching immediately starts working. "
    "When a community discussion sparks, professionals grow their networks. This is the flywheel "
    "effect that makes Profcaria more valuable for every user, every day."
)


# ═══════════════════════════════════════════════════════════
# 8. FOR EMPLOYERS
# ═══════════════════════════════════════════════════════════

doc.add_heading('8. For Employers — Hiring Made Smarter', level=1)
add_divider(doc)

add_body(doc,
    "Profcaria gives HR teams and recruiters a suite of intelligent tools designed to reduce "
    "time-to-hire, improve candidate quality, and lower recruitment costs:"
)

employer_benefits = [
    ("Faster Screening", "AI Top Match analyses every application and ranks candidates by fit — eliminating hours of manual CV review."),
    ("Better Decisions", "Rich candidate profiles with verified badges, references, skills, and portfolio work give recruiters the full picture before the interview."),
    ("Complete Pipeline Visibility", "Track every application from submitted → shortlisted → employed with real-time analytics. Identify bottlenecks and optimise your process."),
    ("Employer Branding", "Your company profile and feed posts attract passive candidates who align with your culture and values."),
    ("Cost Efficiency", "Plans start free and scale affordably. No per-seat charges, no hidden fees, no contracts."),
    ("Local Targeting", "Location-restricted postings ensure you reach candidates where you need them — not across the entire country."),
]

for title, desc in employer_benefits:
    add_bullet(doc, desc, bold_prefix=title)


# ═══════════════════════════════════════════════════════════
# 9. FOR PROFESSIONALS
# ═══════════════════════════════════════════════════════════

doc.add_heading('9. For Professionals — Your Career, Elevated', level=1)
add_divider(doc)

add_body(doc,
    "Profcaria is built for professionals who want more than just a job board — they want a "
    "platform that values their growth:"
)

prof_benefits = [
    ("Skill-Based Discovery", "Get found by employers based on what you can actually do — not just your job title or university."),
    ("Verification Badges", "Stand out in a competitive market with verification badges that boost your visibility by up to 8x."),
    ("Transparent Applications", "Track your application status in real time. Know exactly when you have been shortlisted, and receive instant notifications."),
    ("Professional Portfolio", "Showcase your work, projects, certifications, and references — all in one place that employers will actually see."),
    ("Community Access", "Join industry communities, participate in discussions, and build a network that creates real career opportunities."),
    ("Zero Cost Entry", "Access to jobs, networking, and communities is completely free. Premium features are optional and affordable."),
]

for title, desc in prof_benefits:
    add_bullet(doc, desc, bold_prefix=title)


# ═══════════════════════════════════════════════════════════
# 10. PRICING PLANS
# ═══════════════════════════════════════════════════════════

doc.add_heading('10. Pricing Plans', level=1)
add_divider(doc)

doc.add_heading('Employer Plans', level=2)

add_body(doc, "Simple, transparent pricing. No per-seat charges. No contracts. Cancel anytime.")

employer_plans = [
    ['Free', 'KSh 0/mo', '1', '1 Year', '—', '—'],
    ['Basic', '~KSh 3,225/mo', '5', '3 Years', '5 Credits', '—'],
    ['Pro (Best Value)', '~KSh 9,675/mo', '30', 'Unlimited', '15 Credits (3/Job)', '✅'],
    ['Enterprise', '~KSh 21,930/mo', 'Unlimited', 'Unlimited', 'Unlimited', '✅'],
]

add_styled_table(doc,
    ['Plan', 'Price', 'Job Posts/mo', 'Analytics', 'AI Top Match', 'Location Lock'],
    employer_plans,
    col_widths=[3, 3, 2.5, 2.5, 3, 2.5]
)

doc.add_paragraph()

doc.add_heading('Professional Plans', level=2)

add_body(doc, "Boost your visibility and stand out to employers.")

prof_plans = [
    ['Free', 'KSh 0/mo', '—', 'Standard', 'Standard Feed'],
    ['Basic', '~KSh 645/mo', 'Gray Badge', '1.5x Visibility', 'Appear in Verified Filters'],
    ['Standard', '~KSh 1,935/mo', 'Blue Badge', '3.5x Visibility', 'Stand Out to Employers'],
    ['Premium', '~KSh 3,870/mo', 'Gold Badge', '8x Max Visibility', 'Top of Employer Searches'],
]

add_styled_table(doc,
    ['Plan', 'Price', 'Badge', 'Visibility Boost', 'Key Benefit'],
    prof_plans,
    col_widths=[2.5, 2.5, 2.5, 3, 4.5]
)

doc.add_paragraph()
add_body(doc, "* Prices shown in Kenyan Shillings (KSh). Prices are converted from USD at current exchange rates and may vary slightly.")


# ═══════════════════════════════════════════════════════════
# 11. HOW WE COMPARE
# ═══════════════════════════════════════════════════════════

doc.add_heading('11. How We Compare', level=1)
add_divider(doc)

add_body(doc,
    "Profcaria offers enterprise-grade features at a fraction of the cost of established platforms. "
    "Here is how we stack up:"
)

compare = [
    ['Monthly Cost (Hiring)', '~KSh 9,675 (Pro)', '~KSh 22,000+', '~KSh 13,000–26,000'],
    ['Job Posts Included', '30/month', 'Limited (credit-based)', 'Pay per listing'],
    ['AI Candidate Matching', '✅ 15 credits/mo', '❌ Basic suggestions', '❌ Manual only'],
    ['Hiring Analytics Dashboard', '✅ Full analytics', '❌ Limited reporting', '❌ Basic'],
    ['Location-Restricted Posting', '✅ Built-in', '❌', '❌'],
    ['Skills-First Filtering', '✅ Toggle per job', '❌', '❌'],
    ['Verification Badges', '✅ 3 tiers', '❌', '❌'],
    ['Reference System', '✅ On-platform', '❌', '❌'],
    ['Professional Communities', '✅ Built-in', '✅ Groups', '❌'],
    ['Ads & Distractions', '❌ No ads', '✅ Heavy ads', '✅ Sponsored content'],
    ['Per-Seat Pricing', '❌ One flat price', '✅ Per seat', '✅ Per package'],
    ['Local Payment (M-Pesa, Card)', '✅ Paystack', '❌ International only', '✅ Limited'],
]

add_styled_table(doc,
    ['Feature', 'Profcaria', 'LinkedIn Recruiter', 'Fuzu Employer'],
    compare,
    col_widths=[4, 3.5, 3.5, 3.5]
)


# ═══════════════════════════════════════════════════════════
# 12. NO ADS
# ═══════════════════════════════════════════════════════════

doc.add_heading('12. No Ads. Ever.', level=1)
add_divider(doc)

add_body(doc,
    "One of the core promises of Profcaria is a completely ad-free experience. We will never "
    "show banner ads, sponsored posts, or promotional content in your feed."
)

add_body(doc,
    "Why? Because we believe a professional platform should respect your attention. Your feed "
    "should show you relevant updates from your network and industry — not advertisements from "
    "brands bidding for your eyeballs."
)

add_body(doc,
    "Our business model is subscription-based, which means our incentive is simple: build a "
    "platform so valuable that you choose to pay for premium features. Not a platform that sells "
    "your attention to the highest bidder."
)

add_body(doc,
    "This is a deliberate choice. It means slower growth, but deeper trust. And we believe that "
    "trust is the most important currency in a professional ecosystem."
)


# ═══════════════════════════════════════════════════════════
# 13. SECURITY & TRUST
# ═══════════════════════════════════════════════════════════

doc.add_heading('13. Security & Trust', level=1)
add_divider(doc)

add_body(doc, "Profcaria takes data security and user trust seriously:")

security_items = [
    ("End-to-End Encryption", "All sensitive data — including emails, phone numbers, and personal identifiers — is encrypted at rest and in transit."),
    ("Secure Authentication", "JWT-based session management with secure, HTTP-only cookies. No tokens stored in local storage."),
    ("Activity Logging", "Complete security audit trail. Users can see every login, password change, and account action with timestamps and IP addresses."),
    ("Paystack-Secured Payments", "All payments are processed through Paystack, a PCI-DSS compliant payment processor trusted by thousands of African businesses."),
    ("Verification System", "Multi-tier badge verification ensures that profiles on the platform are genuine, reducing fraud and misrepresentation."),
    ("Link Safety", "Built-in URL safety checking protects users from malicious links in messages and posts."),
]

for title, desc in security_items:
    add_bullet(doc, desc, bold_prefix=title)


# ═══════════════════════════════════════════════════════════
# 14. WHY PROFCARIA
# ═══════════════════════════════════════════════════════════

doc.add_heading('14. Why Profcaria', level=1)
add_divider(doc)

add_body(doc, "In a landscape crowded with job boards and social networks, Profcaria stands apart because:")

why = [
    "We are built for Africa — not adapted from a Western product.",
    "We use AI to make hiring decisions better, not just faster.",
    "We charge fairly — no per-seat pricing, no hidden fees, no expiring credits.",
    "We respect your attention — no ads, no spam, no noise.",
    "We are a full ecosystem — networking, hiring, analytics, communities, and career growth in one place.",
    "We verify professionals — so employers can trust what they see.",
    "We are easy to use — intuitive design that requires zero training.",
    "We are constantly improving — built by people who care about getting it right.",
]

for item in why:
    add_bullet(doc, item)


# ═══════════════════════════════════════════════════════════
# 15. CONTACT
# ═══════════════════════════════════════════════════════════

doc.add_heading('15. Contact Information', level=1)
add_divider(doc)

contact_info = [
    ("Website", "www.profcaria.com"),
    ("Email", "[Your Email]"),
    ("Phone", "[Your Phone]"),
    ("Location", "Nairobi, Kenya"),
]

for label, value in contact_info:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run = p.add_run(value)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(4)

doc.add_paragraph()

closing = doc.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
closing.paragraph_format.space_before = Pt(24)
run = closing.add_run("Where Ambition Finds Its Home.")
run.italic = True
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
run.font.name = 'Calibri'


# ═══════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════

output_path = os.path.join(os.path.dirname(__file__), '..', 'Profcaria_Company_Profile.docx')
output_path = os.path.abspath(output_path)
doc.save(output_path)
print(f"✅ Document saved to: {output_path}")
